# algorithm.py
"""
Implements the closest pair of points algorithm using divide-and-conquer.
All steps are logged and integrated into a Word report, along with relevant figures.
Distances are computed once per pair, no redundant instructions or calculations appear.
"""

from collections import defaultdict
from typing import List
import matplotlib.pyplot as plt
import os
from datetime import datetime
from docx.shared import Inches

from .geometry import ColoredPoint, dist
from .logger import log_message, doc

# Create a unique timestamp and directories
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
base_output_dir = f"output/{timestamp}"
plot_output_dir = f"{base_output_dir}/plot"
array_output_dir = f"{base_output_dir}/array"
os.makedirs(plot_output_dir, exist_ok=True)
os.makedirs(array_output_dir, exist_ok=True)

# Global states
global_step_counter = 0
global_best_dist = float('inf')
global_best_pair = None
global_dividers = []

def group_by_color(points: List[ColoredPoint]):
    groups = defaultdict(list)
    for point in points:
        groups[point.color].append(point)
    return groups

def update_global_best_pair(p1: ColoredPoint, p2: ColoredPoint, d: float):
    global global_best_dist, global_best_pair
    if d < global_best_dist:
        global_best_dist = d
        global_best_pair = (p1, p2, d)
        log_message(f"[update_global_best_pair] Updated global best pair: {p1}, {p2} with distance {d}")

from docx.shared import Inches

def insert_image_to_doc(document, img_path, caption=""):
    document.add_picture(img_path, width=Inches(5))
    if caption:
        para = document.add_paragraph(caption)
        para.alignment = 1

def plot_array(points: List[ColoredPoint],
               title="",
               highlight_index=None,
               left_size=None,
               save_fig=True,
               arrow_index=None):
    global global_step_counter, global_best_dist

    fig, ax = plt.subplots(figsize=(10, 2))
    ax.set_axis_off()
    n = len(points)
    table = plt.table(cellText=[["" for _ in range(n)]],
                      cellLoc='center', loc='center', edges='closed')

    for i in range(n):
        table.auto_set_column_width(i)

    for i, p in enumerate(points):
        cell = table[0, i]
        cell.get_text().set_text(f"{p.x},{p.y}")
        cell.set_facecolor(p.color if p.color else 'white')
        text_color = 'white' if p.color == 'black' else 'black'
        cell.get_text().set_color(text_color)

    if highlight_index is not None and highlight_index < n:
        ax.annotate("↓ divide here",
                    xy=(highlight_index + 1, 1.05), xycoords=('data', 'axes fraction'),
                    xytext=(highlight_index + 1, 1.5), textcoords=('data', 'axes fraction'),
                    arrowprops=dict(facecolor='red', shrink=0.05),
                    ha='center', va='bottom', color='red', fontsize=10)

    if left_size is not None and left_size < n:
        ax.plot([left_size, left_size], [1.2, -0.2], color='red', linestyle='--', linewidth=2, transform=ax.transData)

    if arrow_index is not None and arrow_index < n:
        ax.annotate("← mid", xy=(arrow_index + 0.5, 0.5), xycoords=('data', 'axes fraction'),
                    xytext=(arrow_index + 0.5, 1.2), textcoords=('data', 'axes fraction'),
                    arrowprops=dict(facecolor='blue', shrink=0.05),
                    ha='center', va='bottom', color='blue', fontsize=10)

    if global_best_dist < float('inf'):
        plt.title(f"{title}\nGlobal minimum δ: {global_best_dist:.4f}")
    else:
        plt.title(title)

    plt.tight_layout()
    if save_fig:
        filename = f"{array_output_dir}/closest_pair_array_step_{global_step_counter}.png"
        plt.savefig(filename, dpi=150)

        doc.add_heading(title, level=2)
        insert_image_to_doc(doc, filename, caption=title)
        doc.add_page_break()
    else:
        plt.show()
    plt.close()

def plot_points(original_points: List[ColoredPoint],
                title="",
                vertical_line_x=None,
                pairs=[],  # list of (p1, p2, d_ij)
                strip_line_x=None,
                δ=None,
                highlight_points=None,
                save_fig=True,
                divide_label=None):
    global global_step_counter, global_best_dist, global_dividers

    plt.figure(figsize=(8, 8))
    groups = group_by_color(original_points)
    for color, pts in groups.items():
        xs = [p.x for p in pts]
        ys = [p.y for p in pts]
        plt.scatter(xs, ys, c=color, s=30, edgecolors='black', linewidths=0.5, zorder=2, alpha=0.3,
                    label=f"{color} points")

    if highlight_points:
        h_groups = group_by_color(highlight_points)
        for c, hpts in h_groups.items():
            hx = [hp.x for hp in hpts]
            hy = [hp.y for hp in hpts]
            plt.scatter(hx, hy, c=c, s=100, edgecolors='black', linewidths=1, zorder=3, alpha=1.0,
                        label='Highlighted Points')

    if pairs:
        compared_points = set()
        for (p1, p2, d_ij) in pairs:
            plt.plot([p1.x, p2.x], [p1.y, p2.y], 'g--', linewidth=1.5, zorder=4, alpha=1.0,
                     label='Comparison Pair')
            mid_x = (p1.x + p2.x) / 2
            mid_y = (p1.y + p2.y) / 2
            plt.text(mid_x, mid_y, f"δ={d_ij:.4f}", color='black', fontsize=10, ha='center', va='bottom')
            compared_points.add(p1)
            compared_points.add(p2)

        c_groups = defaultdict(list)
        for cp in compared_points:
            c_groups[cp.color].append(cp)
        for c, cpts in c_groups.items():
            cx = [cp.x for cp in cpts]
            cy = [cp.y for cp in cpts]
            plt.scatter(cx, cy, c=c, s=100, edgecolors='black', linewidths=1, zorder=5, alpha=1.0,
                        label='Compared Points')

    if global_best_pair is not None:
        p1, p2, best_dist = global_best_pair
        plt.plot([p1.x, p2.x], [p1.y, p2.y], 'r-', linewidth=3.0, zorder=6, alpha=1.0, label='Global Minimum Pair')
        mid_x = (p1.x + p2.x)/2
        mid_y = (p1.y + p2.y)/2
        plt.text(mid_x, mid_y, f"Global Min δ={best_dist:.4f}", color='red', fontsize=10, ha='center', va='top')

    for (dx, dlabel) in global_dividers:
        plt.axvline(x=dx, color='red', linestyle='--', linewidth=1.5, zorder=1, alpha=0.3)
        if dlabel is not None:
            plt.text(dx, 64, f"$n_{{{dlabel}}}$", color='red', fontsize=12, ha='center', va='top', alpha=0.3)

    if vertical_line_x is not None:
        plt.axvline(x=vertical_line_x, color='red', linestyle='--', linewidth=1.5,
                    label='Dividing Line', zorder=1, alpha=1.0)
        if divide_label is not None:
            plt.text(vertical_line_x, 64, f"$n_{{{divide_label}}}$", color='red', fontsize=12, ha='center',
                     va='top', alpha=1.0)

    if strip_line_x is not None:
        left_x, right_x = strip_line_x
        plt.axvspan(left_x, right_x, color='yellow', alpha=0.2, label='Strip Region', zorder=0)

    if δ is not None:
        plt.title(f"{title}\nLocal minimal δ: {δ:.4f}\nGlobal minimal δ: {global_best_dist:.4f}")
    else:
        if global_best_dist < float('inf'):
            plt.title(f"{title}\nGlobal minimal δ: {global_best_dist:.4f}")
        else:
            plt.title(title)

    plt.xlim(0, 64)
    plt.ylim(0, 64)
    plt.xticks(range(0, 61, 10))
    plt.yticks(range(0, 61, 10))
    plt.grid(which='major', color='black', linestyle='-', linewidth=1)
    plt.minorticks_on()
    plt.grid(which='minor', color='grey', linestyle=':', linewidth=0.5)

    handles, labels = plt.gca().get_legend_handles_labels()
    by_label = dict(zip(labels, handles))
    plt.legend(by_label.values(), by_label.keys(), loc='best')

    plt.xlabel('X')
    plt.ylabel('Y')
    plt.tight_layout()

    if save_fig:
        filename = f"{plot_output_dir}/closest_pair_step_{global_step_counter}.png"
        plt.savefig(filename, dpi=150)

        doc.add_heading(title, level=2)
        insert_image_to_doc(doc, filename, caption=title)
        doc.add_page_break()

        global_step_counter += 1
    else:
        plt.show()

    plt.close()

def brute_force(original_points: List[ColoredPoint], subset_points: List[ColoredPoint]) -> float:
    log_message("====================================================================================================")
    log_message("[brute_force] Brute forcing the closest pair in a small subset of points.")
    for p in subset_points:
        log_message(f"             {p}")
    n = len(subset_points)
    δ = float("inf")

    plot_points(original_points, title="Brute Force on Small Subset",
                highlight_points=subset_points, δ=δ, pairs=[])
    plot_array(subset_points, title="Brute Force Subset Array")

    for i in range(n):
        for j in range(i + 1, n):
            log_message("---------------------------------------------------------------------------------------------------")
            log_message(f"[brute_force] Checking pair: {subset_points[i]} and {subset_points[j]}")
            d_ij = dist(subset_points[i], subset_points[j])
            log_message(f"[brute_force] Distance = {d_ij}")

            pairs_with_dist = [(subset_points[i], subset_points[j], d_ij)]

            plot_points(original_points, title="Brute Force Comparison",
                        pairs=pairs_with_dist, highlight_points=subset_points, δ=δ)
            plot_array(subset_points, title="Brute Force Comparison Array")

            if d_ij < δ:
                log_message(f"[brute_force] Found a smaller local distance: {d_ij} < {δ}")
                δ = d_ij
                log_message(f"[brute_force] Updating local δ = {δ}")
                update_global_best_pair(subset_points[i], subset_points[j], δ)
                pairs_with_dist = [(subset_points[i], subset_points[j], δ)]
                plot_points(original_points, title="Brute Force New Minimum Found",
                            pairs=pairs_with_dist, highlight_points=subset_points, δ=δ)
                plot_array(subset_points, title="Brute Force New Min Array")
            else:
                plot_points(original_points, title="Brute Force No Improvement",
                            pairs=pairs_with_dist, highlight_points=subset_points, δ=δ)
                plot_array(subset_points, title="Brute Force No Improvement Array")

    log_message(f"[brute_force] Brute force complete. Minimum distance found (locally): {δ}")
    return δ

def strip_closest(original_points: List[ColoredPoint], strip: List[ColoredPoint], δ: float) -> float:
    log_message("====================================================================================================")
    log_message("[strip_closest] Checking the 'strip' area for possibly closer pairs.")
    log_message(f"[strip_closest] Current δ: {δ}")
    for p in strip:
        log_message(f"                {p}")

    strip.sort(key=lambda p: p.y)
    log_message("[strip_closest] Points in strip after sorting by y-coordinate:")
    for p in strip:
        log_message(f"                {p}")

    min_dist = δ
    n = len(strip)

    if len(strip) > 0:
        plot_points(original_points, title="Checking Strip for Closer Pairs", δ=min_dist, pairs=[])
        plot_array(strip, title="Strip Array (after sorting by y)")

    for i in range(n):
        for j in range(i + 1, min(i + 7, n)):
            vertical_distance = strip[j].y - strip[i].y
            if vertical_distance >= min_dist:
                log_message("----------------------------------------------------------------------------------------------------")
                log_message(f"[strip_closest] For {strip[i]} and {strip[j]}, vertical dist {vertical_distance} >= {min_dist}. Stop inner loop.")
                break
            else:
                log_message("---------------------------------------------------------------------------------------------------")
                log_message(f"[strip_closest] Checking {strip[i]} vs {strip[j]}")
                log_message(f"[strip_closest] Vertical distance {vertical_distance} < {min_dist}, calculating distance.")
                d_ij = dist(strip[i], strip[j])
                log_message(f"[strip_closest] Computed distance = {d_ij}")

                pairs_with_dist = [(strip[i], strip[j], d_ij)]
                plot_points(original_points, title="Strip Comparison",
                            pairs=pairs_with_dist, δ=min_dist)
                plot_array(strip, title="Strip Comparison Array")

                if d_ij < min_dist:
                    log_message(f"[strip_closest] Found smaller local distance in strip: {d_ij} < {min_dist}")
                    min_dist = d_ij
                    update_global_best_pair(strip[i], strip[j], min_dist)
                    pairs_with_dist = [(strip[i], strip[j], d_ij)]
                    plot_points(original_points, title="Strip New Minimum Found",
                                pairs=pairs_with_dist, δ=min_dist)
                    plot_array(strip, title="Strip New Minimum Array")
                else:
                    plot_points(original_points, title="Strip Comparison No Improvement",
                                pairs=pairs_with_dist, δ=min_dist)
                    plot_array(strip, title="Strip No Improvement Array")

    log_message(f"[strip_closest] Strip check complete. Minimum distance in strip: {min_dist}")
    return min_dist

def closest_pair_util(original_points: List[ColoredPoint], subset_points: List[ColoredPoint], depth=0) -> float:
    global global_dividers
    log_message("====================================================================================================")
    log_message("[closest_pair_util] Divide and Conquer step")
    log_message(f"[closest_pair_util] Number of points in subset: {len(subset_points)}")
    for p in subset_points:
        log_message(f"                 {p}")

    plot_array(subset_points, title=f"Array at depth {depth}")

    n = len(subset_points)
    if n <= 3:
        log_message("[closest_pair_util] n <= 3, using brute force.")
        return brute_force(original_points, subset_points)

    mid = n // 2
    mid_point = subset_points[mid]
    log_message("[closest_pair_util] Dividing the subset into two halves.")
    log_message(f"[closest_pair_util] Mid index: {mid}, Mid point: {mid_point}")
    left_points = subset_points[:mid]
    right_points = subset_points[mid:]

    global_dividers.append((mid_point.x, depth))

    plot_points(original_points, title="Divide Step", vertical_line_x=mid_point.x, δ=float('inf'), divide_label=depth, pairs=[])
    plot_array(subset_points, title=f"Divide Step Array (depth {depth})", left_size=len(left_points), arrow_index=mid)

    log_message("[closest_pair_util] Left subset points:")
    for p in left_points:
        log_message(f"                  {p}")
    log_message("[closest_pair_util] Right subset points:")
    for p in right_points:
        log_message(f"                  {p}")

    log_message("[closest_pair_util] Recursively solving left half:")
    δ_left = closest_pair_util(original_points, left_points, depth=depth + 1)
    log_message(f"[closest_pair_util] Closest pair distance in left half: {δ_left}")

    log_message("[closest_pair_util] Recursively solving right half:")
    δ_right = closest_pair_util(original_points, right_points, depth=depth + 1)
    log_message(f"[closest_pair_util] Closest pair distance in right half: {δ_right}")

    δ = min(δ_left, δ_right)
    log_message(f"[closest_pair_util] Minimum distance between left and right subsets: {δ}")

    log_message("[closest_pair_util] Building strip of points close to dividing line.")
    strip = [p for p in subset_points if abs(p.x - mid_point.x) < δ]
    log_message("[closest_pair_util] Points in strip (|x - mid_point.x| < δ):")
    for p in strip:
        log_message(f"                  {p}")

    if len(strip) > 0:
        strip_left = mid_point.x - δ
        strip_right = mid_point.x + δ
        plot_points(original_points, title="Strip Construction", strip_line_x=(strip_left, strip_right), δ=δ, pairs=[])
        plot_array(strip, title=f"Strip Array (depth {depth})")

    δ_strip = strip_closest(original_points, strip, δ)
    log_message(f"[closest_pair_util] Minimum distance found in strip: {δ_strip}")

    final_min = min(δ, δ_strip)
    log_message(f"[closest_pair_util] Returning minimum distance from this subproblem: {final_min}")

    plot_points(original_points, title="Subproblem Result", δ=final_min, pairs=[])
    plot_array(subset_points, title=f"Subproblem Result Array (depth {depth})")

    return final_min

def closest_pair_distance(points: List[ColoredPoint]) -> float:
    global global_best_dist, global_best_pair, global_dividers, global_step_counter
    global_best_dist = float('inf')
    global_best_pair = None
    global_dividers = []
    global_step_counter = 0

    # Modern style: add a cover page, introduction, etc.
    doc.add_heading('Closest Pair of Points Analysis', 0)
    doc.add_paragraph("This report presents a step-by-step analysis of the Closest Pair of Points problem, utilizing a divide-and-conquer strategy. All intermediate steps, computations, and visualizations are included. The global minimum distance is highlighted, and the final results are summarized at the end.").alignment = 0
    doc.add_page_break()

    log_message("====================================================================================================")
    log_message("[closest_pair_distance] Initiating closest pair computation.")
    if not points or len(points) < 2:
        log_message("[closest_pair_distance] Not enough points provided. Returning inf.")
        return float('inf')

    log_message("[closest_pair_distance] Sorting points by x-coordinate.")
    points_sorted = sorted(points, key=lambda p: p.x)
    log_message("[closest_pair_distance] Points after sorting by x:")
    for p in points_sorted:
        log_message(f"                       {p}")

    plot_points(points_sorted, title="Initial Set of Points", pairs=[])
    plot_array(points_sorted, title="Initial Sorted Array")

    result = closest_pair_util(points_sorted, points_sorted, depth=0)
    log_message(f"[closest_pair_distance] Closest pair distance result: {result}")

    plot_points(points_sorted, title="Final Result", δ=result, pairs=[])
    plot_array(points_sorted, title="Final Array Result")

    doc.add_heading('Analysis Completed', level=1)
    doc.add_paragraph(f"The closest pair distance found: {result:.4f}")
    doc.add_paragraph("All computations, intermediate steps, and figures are presented above, providing a transparent overview of the algorithm's process.")

    doc_path = f"{base_output_dir}/report.docx"
    doc.save(doc_path)
    log_message(f"[closest_pair_distance] Report saved at {doc_path}")

    return result
